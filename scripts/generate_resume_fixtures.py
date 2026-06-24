"""Generate the four synthetic resume fixtures committed under tests/fixtures/.

Phase 7 D-Discretion / T-07-foundation-PII threat: all content is SYNTHETIC.
No fragment of Adrian's real resume content appears in any fixture. Every
text-bearing fixture includes a "TEST FIXTURE -- synthetic data" watermark.

Fixtures produced:
- tests/fixtures/sample-resume.pdf       -- valid PDF, extracts >=100 chars
- tests/fixtures/sample-resume.docx      -- valid DOCX with paragraphs + a table
- tests/fixtures/encrypted-sample.pdf    -- password-protected PDF (user_password='test')
- tests/fixtures/empty-text-sample.pdf   -- blank-page PDF, extracts <100 chars

Run from repo root:
    uv run python scripts/generate_resume_fixtures.py

This script is a one-shot helper, NOT production code. It lives outside src/
so it cannot accidentally be imported by the application.
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import (
    ArrayObject,
    DecodedStreamObject,
    DictionaryObject,
    FloatObject,
    NameObject,
    NumberObject,
)

FIXTURE_DIR = Path("tests/fixtures")

SAMPLE_RESUME_TEXT_LINES = [
    "TEST FIXTURE -- synthetic data",
    "Jane Doe -- AI Engineer",
    "Skills: Python, FastAPI, PostgreSQL, pgvector, Docker, Azure Container Apps",
    "Languages: English, German",
    "Target roles: AI Engineer, ML Engineer",
    "Min salary EUR: 70000",
    "Years experience: 5",
]


def _build_text_pdf_bytes(lines: list[str]) -> bytes:
    """Return PDF bytes containing one page that renders the given text lines.

    Builds a minimal page content stream with Tj operators against the Helvetica
    type-1 font. Uses pypdf only -- no reportlab dependency. The resulting PDF
    is valid enough for pypdf.PdfReader.extract_text() to return the original
    text concatenated with newlines.
    """
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    # Build a content stream: BT ... ET with Tf, Td, Tj per line.
    content_parts: list[str] = ["BT", "/F1 12 Tf"]
    # Start near the top; move down 16pt per line.
    y = 750
    content_parts.append(f"50 {y} Td")
    for i, line in enumerate(lines):
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        if i == 0:
            content_parts.append(f"({escaped}) Tj")
        else:
            content_parts.append("0 -16 Td")
            content_parts.append(f"({escaped}) Tj")
    content_parts.append("ET")
    content_stream = "\n".join(content_parts).encode("latin-1")

    stream_obj = DecodedStreamObject()
    stream_obj.set_data(content_stream)
    page[NameObject("/Contents")] = writer._add_object(stream_obj)  # noqa: SLF001

    # Register a Helvetica /F1 font so the /F1 Tf operator resolves.
    font_dict = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
            NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
        }
    )
    font_ref = writer._add_object(font_dict)  # noqa: SLF001
    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref}),
        }
    )
    page[NameObject("/Resources")] = resources
    page[NameObject("/MediaBox")] = ArrayObject(
        [NumberObject(0), NumberObject(0), FloatObject(612), FloatObject(792)]
    )

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def write_sample_resume_pdf() -> Path:
    path = FIXTURE_DIR / "sample-resume.pdf"
    data = _build_text_pdf_bytes(SAMPLE_RESUME_TEXT_LINES)
    path.write_bytes(data)
    # Self-check: extract text and confirm >=100 non-whitespace chars
    reader = PdfReader(io.BytesIO(data))
    extracted = "\n".join(p.extract_text() or "" for p in reader.pages)
    assert reader.is_encrypted is False, "sample PDF must not be encrypted"
    assert len(extracted.strip()) >= 100, (
        f"sample PDF text too short: {len(extracted.strip())} chars\n{extracted!r}"
    )
    return path


def write_sample_resume_docx() -> Path:
    path = FIXTURE_DIR / "sample-resume.docx"
    doc = Document()
    doc.add_heading("TEST FIXTURE -- synthetic data", level=1)
    doc.add_paragraph("Jane Doe -- AI Engineer")
    doc.add_paragraph(
        "Skills: Python, FastAPI, PostgreSQL, pgvector, Docker, Azure Container Apps"
    )
    doc.add_paragraph("Languages: English, German")
    doc.add_paragraph("Target roles: AI Engineer, ML Engineer")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Years"
    table.rows[0].cells[1].text = "5"
    table.rows[1].cells[0].text = "Min salary (EUR)"
    table.rows[1].cells[1].text = "70000"
    doc.save(path)
    return path


def write_encrypted_sample_pdf() -> Path:
    path = FIXTURE_DIR / "encrypted-sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.encrypt(user_password="test", owner_password="test")
    with open(path, "wb") as f:
        writer.write(f)
    # Self-check
    with open(path, "rb") as f:
        reader = PdfReader(f)
        assert reader.is_encrypted, "encrypted fixture must report is_encrypted=True"
    return path


def write_empty_text_sample_pdf() -> Path:
    path = FIXTURE_DIR / "empty-text-sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with open(path, "wb") as f:
        writer.write(f)
    # Self-check: <100 chars on extract
    with open(path, "rb") as f:
        reader = PdfReader(f)
        extracted = "\n".join(p.extract_text() or "" for p in reader.pages)
    assert len(extracted.strip()) < 100, (
        f"empty-text fixture extracts too much: {len(extracted.strip())} chars"
    )
    return path


def main() -> None:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    paths = [
        write_sample_resume_pdf(),
        write_sample_resume_docx(),
        write_encrypted_sample_pdf(),
        write_empty_text_sample_pdf(),
    ]
    for p in paths:
        size = p.stat().st_size
        print(f"  wrote {p} ({size} bytes)")
    print("OK")


if __name__ == "__main__":
    main()
